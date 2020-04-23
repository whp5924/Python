#

from selenium import webdriver
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH #   用来居中显示标题
from time import sleep
from selenium.webdriver.common.keys import Keys
'''

'''
# 浏览器安装路径
#BROWSER_PATH=\'C:\Users\Administrator\AppData\Local\Google\Chrome\Application\chromedriver.exe'
#目的URL
DEST_URL='https://wenku.baidu.com/view/aa31a84bcf84b9d528ea7a2c.html'
#用来保存文档
doc_title = ''
doc_content_list = []
def find_doc(driver, init=True):
  global doc_content_list
  global doc_title
  stop_condition = False
  html = driver.page_source
  soup1 = BeautifulSoup(html, 'lxml')
  if (init is True): # 得到标题
    title_result = soup1.find('div', attrs={'class': 'doc-title'})
    #doc_title = title_result.get_text() # 得到文档标题
    doc_title = title_result.get

    # 拖动滚动条
    init_page = driver.find_element_by_xpath( "//div[@class='foldpagewg-text-con']")
    print(type(init_page), init_page)
    driver.execute_script('arguments[0].scrollIntoView();', init_page)
    init_page.click()
    init = False
  else:
    try:
      page = driver.find_element_by_xpath( "//div[@class='pagerwg-schedule']")
      #print(type(next_page), next_page)
      next_page = driver.find_element_by_class_name("pagerwg-button")
      station = driver.find_element_by_xpath( "//div[@class='bottombarwg-root border-none']")
      driver.execute_script('arguments[0].scrollIntoView(false);', station)
      #js.executeScript("arguments[0].click();",next_page);
      #sleep(5)
      '''js = "window.scrollTo(508,600)"
      driver.execute_script(js)'''
      next_page.click()
    except:
      #结束条件
      print("找不到元素")
      stop_condition = True
      #next_page.send_keys(Keys.ENTER)
      # 遍历所有的txt标签标定的文档，将其空格删除，然后进行保存
  content_result = soup1.find_all('p', attrs={'class': 'txt'})
  for each in content_result:
    each_text = each.get_text()
    if ' ' in each_text:
      text = each_text.replace(' ', '')
    else:
      text = each_text
    # print(each_text)
    doc_content_list.append(text)
          # 得到正文内容
  sleep(2) # 防止页面加载过慢
  if stop_condition is False:
    doc_title, doc_content_list = find_doc(driver, init)
  return doc_title, doc_content_list
def save(doc_title, doc_content_list):
  document = Document()
  heading = document.add_heading(doc_title, 0)
  heading.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中显示
  for each in doc_content_list:
    document.add_paragraph(each)
  # 处理字符编码问题
  t_title = doc_title.split()[0]
  #print(t_title)
  #document.save('2.docx')
  document.save('百度文库-%s.docx'% t_title)
  print("\n\nCompleted: %s.docx, to read." % t_title)
  driver.quit()
if __name__ == '__main__':
  options = webdriver.ChromeOptions()
  options.add_argument('user-agent="Mozilla/5.0 (Linux; Android 4.0.4; \ Galaxy Nexus Build/IMM76B) AppleWebKit/535.19 (KHTML, like Gecko) \ Chrome/18.0.1025.133 Mobile Safari/535.19"')
  #driver = webdriver.Chrome(BROWSER_PATH, chrome_options=options)
  driver = webdriver.Chrome(chrome_options=options)
  driver.get(DEST_URL)
  #JavascriptExecutor js = (JavascriptExecutor) driver;
  print("**********START**********")
  title, content = find_doc(driver, True)
  save(title, content)
  driver.quit()


